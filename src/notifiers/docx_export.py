"""Экспорт карточки рекомендации в Word (.docx)."""

from __future__ import annotations

import re
from datetime import date
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from src.forecast.recommendation_instructions import (
    STATUS_LABELS,
    TYPE_LABELS,
    format_date_ru,
    format_price_rub,
)


def _slug_category(label: str) -> str:
    text = label.lower().strip()
    text = re.sub(r"[^\w\s\-а-яё]+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"[\s_]+", "-", text)
    return text[:60] or "category"


def recommendation_docx_filename(
    rec_id: int,
    target_date: date,
    room_label: str,
) -> str:
    return (
        f"1apart_рекомендация_{rec_id}_"
        f"{target_date.isoformat()}_{_slug_category(room_label)}.docx"
    )


def build_recommendation_docx(card: dict[str, Any]) -> bytes:
    """Собрать .docx по данным карточки (без PII и служебных ключей)."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading("1apart — карточка внедрения рекомендации", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    created = card.get("exported_at") or date.today().isoformat()
    doc.add_paragraph(f"Дата создания документа: {created}")

    if card.get("low_confidence"):
        warn = doc.add_paragraph()
        run = warn.add_run(
            "⚠ НИЗКАЯ ДОСТОВЕРНОСТЬ ПРОГНОЗА — применяйте с осторожностью, "
            "предпочтителен повторный расчёт."
        )
        run.bold = True
        run.font.color.rgb = RGBColor(0xB3, 0x26, 0x1E)

    if card.get("snapshot_missing"):
        note = doc.add_paragraph()
        run = note.add_run(
            "Снимок на момент создания отсутствует — часть цифр актуальна "
            "на момент экспорта."
        )
        run.italic = True

    # 1. Сводка решения
    doc.add_heading("1. Сводка решения", level=1)
    decision = card["decision"]
    lines = [
        f"Рекомендация №{decision['id']}",
        f"Дата проживания: {decision['date_label']}",
        f"Категория: {decision['room_label']}",
        f"Действие: {decision['action_label']}",
        f"Текущая цена: {format_price_rub(decision.get('current_price'))}",
    ]
    if decision.get("rec_min") is not None and decision.get("rec_max") is not None:
        lines.append(
            f"Рекомендованный диапазон: "
            f"{format_price_rub(decision['rec_min'])}–"
            f"{format_price_rub(decision['rec_max'])}"
        )
    if decision.get("show_selected_price") and decision.get("selected_price") is not None:
        lines.append(
            f"Рекомендованная цена к применению: "
            f"{format_price_rub(decision['selected_price'])}"
        )
    elif not decision.get("show_selected_price"):
        lines.append(
            "Рекомендованная цена к применению: не задана "
            "(требуется ручная проверка)"
        )
    lines.append(f"Статус: {decision['status_label']}")
    for line in lines:
        doc.add_paragraph(line)

    # 2. Обоснование
    doc.add_heading("2. Обоснование", level=1)
    why = card.get("why") or {}
    for item in why.get("bullets") or []:
        doc.add_paragraph(item, style="List Bullet")

    # 3. Рынок
    doc.add_heading("3. Сравнение с рынком", level=1)
    market = why.get("market") or {}
    doc.add_paragraph(
        f"Текущая цена: {format_price_rub(market.get('current_price'))}"
    )
    doc.add_paragraph(
        f"Медиана конкурентов: {format_price_rub(market.get('market_median'))}"
    )
    gap = market.get("market_gap_pct")
    if gap is not None:
        doc.add_paragraph(f"Отклонение от рынка: {gap:+.1f}%")
    else:
        doc.add_paragraph("Отклонение от рынка: нет данных")

    # 4. События
    doc.add_heading("4. Влияющие события Томска", level=1)
    events = why.get("events") or []
    if events:
        for ev in events:
            title_ev = ev.get("title") or "Событие"
            impact = ev.get("impact_score")
            dates = ev.get("start_at") or ""
            if ev.get("end_at") and ev["end_at"] != ev.get("start_at"):
                dates += f"–{ev['end_at']}"
            impact_s = f", impact {impact:.0f}" if impact is not None else ""
            doc.add_paragraph(f"{title_ev} ({dates}{impact_s})", style="List Bullet")
    else:
        doc.add_paragraph("Подтверждённых событий на дату нет.")

    # 5. Инструкция
    doc.add_heading("5. Пошаговая инструкция в TravelLine", level=1)
    for i, step in enumerate(card.get("steps") or [], start=1):
        doc.add_paragraph(f"{i}. {step}")

    # 6. Контроль
    doc.add_heading("6. Контрольные точки и откат", level=1)
    control = card.get("control") or {}
    for key in ("check_text", "goal_text", "rollback_text"):
        if control.get(key):
            doc.add_paragraph(control[key])

    # 7. Подпись
    doc.add_heading("7. Отметка менеджера", level=1)
    doc.add_paragraph("ФИО / должность: _______________________________")
    doc.add_paragraph("Дата применения: _______________________________")
    doc.add_paragraph("Фактическая цена: _______________________________")
    doc.add_paragraph("Комментарий: ___________________________________")
    if card.get("manager_comment"):
        doc.add_paragraph(f"Комментарий в системе: {card['manager_comment']}")

    # 8. Дисклеймер
    doc.add_heading("8. Дисклеймер", level=1)
    disc = doc.add_paragraph()
    run = disc.add_run(
        "Рекомендация требует ручного подтверждения, "
        "автоматическая смена цены не выполнялась."
    )
    run.italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_universal_recommendation_docx(card: dict[str, Any]) -> bytes:
    """Word-инструкция для любой рекомендации Центра."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading("1apart — рабочая инструкция", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        f"Дата формирования: {card.get('exported_at') or date.today().strftime('%d.%m.%Y')}"
    )

    doc.add_heading("1. Название и номер", level=1)
    doc.add_paragraph(f"Рекомендация №{card.get('id')}: {card.get('title')}")
    doc.add_paragraph(
        f"Раздел: {card.get('module_label')} · Приоритет: {card.get('priority_label')} · "
        f"Статус: {card.get('status_label')}"
    )

    doc.add_heading("2. Цель и приоритет", level=1)
    doc.add_paragraph(card.get("goal") or "—")
    doc.add_paragraph(card.get("goal_detail") or "")

    doc.add_heading("3. Что происходит (основание)", level=1)
    for line in card.get("what_happens") or []:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("4. Пошаговый план", level=1)
    for i, step in enumerate(card.get("steps") or [], start=1):
        doc.add_paragraph(f"{i}. {step}")

    doc.add_heading("5. Срок и ответственный", level=1)
    doc.add_paragraph(f"Ответственный: {card.get('owner') or '—'}")
    doc.add_paragraph(f"Срок: {card.get('due_at') or card.get('due_hint') or '—'}")

    doc.add_heading("6. Критерии результата", level=1)
    doc.add_paragraph(card.get("check_text") or "")
    for item in card.get("success_criteria") or []:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("7. Если результата нет", level=1)
    for item in card.get("rollback_steps") or []:
        doc.add_paragraph(item, style="List Bullet")
    if card.get("escalation"):
        doc.add_paragraph(f"Эскалация: {card['escalation']}")

    doc.add_heading("8. Отметка о выполнении", level=1)
    doc.add_paragraph("ФИО / должность: _______________________________")
    doc.add_paragraph("Дата выполнения: _______________________________")
    doc.add_paragraph("Комментарий: ___________________________________")

    disc = doc.add_paragraph()
    run = disc.add_run(
        "Рекомендация требует ручного подтверждения; "
        "автоматическое выполнение системой не выполнялось."
    )
    run.italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_recommendations_list_docx(
    rows: list[dict[str, Any]],
    *,
    bucket_label: str = "Все",
) -> bytes:
    """Сводный Word по текущему фильтру списка Центра."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading("1apart — список рекомендаций", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Фильтр: {bucket_label}")
    doc.add_paragraph(f"Дата: {date.today().strftime('%d.%m.%Y')}")
    doc.add_paragraph(f"Всего: {len(rows)}")

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ("№", "Заголовок", "Раздел", "Приоритет", "Статус", "Срок")
    for i, name in enumerate(headers):
        hdr[i].text = name

    for r in rows:
        cells = table.add_row().cells
        cells[0].text = str(r.get("id") or "")
        cells[1].text = str(r.get("title") or "")
        cells[2].text = str(r.get("module_label") or "")
        cells[3].text = str(r.get("priority_label") or "")
        cells[4].text = str(r.get("status_label") or "")
        cells[5].text = str(r.get("due_at") or "—")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_recommendations_list_csv(rows: list[dict[str, Any]]) -> bytes:
    """CSV для открытия в Excel (UTF-8 BOM)."""
    import csv
    from io import StringIO

    buf = StringIO()
    writer = csv.writer(buf, delimiter=";")
    writer.writerow(
        ["id", "title", "module", "priority", "status", "due_at", "owner", "expected_result"]
    )
    for r in rows:
        writer.writerow(
            [
                r.get("id") or "",
                r.get("title") or "",
                r.get("module_label") or "",
                r.get("priority_label") or "",
                r.get("status_label") or "",
                r.get("due_at") or "",
                r.get("owner") or "",
                r.get("expected_result") or "",
            ]
        )
    return ("\ufeff" + buf.getvalue()).encode("utf-8")


def universal_docx_filename(rec_id: int, title: str) -> str:
    slug = _slug_category(title)
    return f"1apart_рекомендация_{rec_id}_{slug}.docx"


# re-export for callers
__all__ = [
    "STATUS_LABELS",
    "TYPE_LABELS",
    "build_recommendation_docx",
    "build_recommendations_list_csv",
    "build_recommendations_list_docx",
    "build_universal_recommendation_docx",
    "format_date_ru",
    "recommendation_docx_filename",
    "universal_docx_filename",
]
